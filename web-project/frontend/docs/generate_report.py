# -*- coding: utf-8 -*-
"""팀플 발표용 프로젝트 평가 보고서 Word 문서 생성"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = r"c:\Users\pc\OneDrive\Desktop\커서\web-app-3\web-project\frontend\docs\학식당_지도_서비스_평가보고서.docx"


def set_korean_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_korean_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True)
    return p


def add_para(doc, text, bold=False, size=11, align=None, space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_korean_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.35
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_korean_font(run, size=10.5)
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_korean_font(run, size=10.5, color=RGBColor(0x33, 0x33, 0x33))
    run.italic = True
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_korean_font(run, size=10, bold=True)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    set_korean_font(run, size=10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 표지
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("계명대학교 학식당 지도 서비스\n")
    set_korean_font(r, size=22, bold=True, color=RGBColor(0x1D, 0x4E, 0xD8))
    r2 = t.add_run("프로젝트 코드 평가 및 발표용 설명서")
    set_korean_font(r2, size=16, bold=True)

    doc.add_paragraph()
    for line in [
        "과목: (팀플 과목명 기입)",
        "팀명: (팀명 기입)",
        "작성일: 2026년 5월 30일",
        "대상 파일: Map.jsx, Map_style.css, mapUtils.js, mobile_map.jsx, mobile_map_style.css",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_korean_font(p.add_run(line), size=11)

    doc.add_page_break()

    # 1. 개요
    add_heading(doc, "1. 프로젝트 개요", 1)
    add_para(
        doc,
        "본 프로젝트는 계명대학교 캠퍼스 내 4개 학식당(공대 식당, 구바, 신바, 아람관)을 지도 위에 표시하고, "
        "사용자의 현재 GPS 위치에서 선택한 식당까지 보행 경로·소요 시간·거리를 안내하는 웹 애플리케이션이다.",
    )
    add_para(doc, "한 줄 요약", bold=True, size=11)
    add_quote(
        doc,
        "GPS로 내 위치를 잡고, 네이버 지도에 학식당을 표시한 뒤, T-map API로 걸어가는 길을 그려주는 반응형 웹 지도 서비스",
    )

    add_para(doc, "기술 스택", bold=True)
    add_bullet(doc, "프론트엔드: React (Create React App)")
    add_bullet(doc, "지도 표시: 네이버 지도 JavaScript API v3")
    add_bullet(doc, "경로 탐색: SK T-map Open API (보행자 경로)")
    add_bullet(doc, "위치 정보: 브라우저 Geolocation API")
    add_bullet(doc, "반응형: 화면 너비 768px 기준 PC/모바일 UI 자동 전환")

    # 2. 아키텍처
    add_heading(doc, "2. 시스템 아키텍처", 1)
    add_para(doc, "애플리케이션은 진입점(index.js)에서 화면 크기에 따라 PC용 Map.jsx 또는 모바일용 mobile_map.jsx를 렌더링한다. 두 컴포넌트는 mapUtils.js의 공통 로직을 공유한다.")
    add_table(
        doc,
        ["파일", "역할", "규모(약)"],
        [
            ["index.js", "768px 기준 PC/모바일 컴포넌트 분기", "31줄"],
            ["Map.jsx", "데스크톱 UI + 지도 상태 관리", "235줄"],
            ["mobile_map.jsx", "모바일 UI + 지도 상태 관리", "213줄"],
            ["mapUtils.js", "식당 좌표, 마커, T-map 경로 API", "94줄"],
            ["Map_style.css", "PC 레이아웃 + 지도 오버레이 스타일", "240줄"],
            ["mobile_map_style.css", "모바일 전용 레이아웃", "128줄"],
        ],
    )

    add_heading(doc, "2.1 데이터 흐름", 2)
    steps = [
        "위치 획득: navigator.geolocation.getCurrentPosition()으로 GPS 수신. 실패 시 DEFAULT_LOCATION(계명대 근처)으로 대체.",
        "지도 초기화: REACT_APP_NAVER_MAP_CLIENT_ID 환경 변수로 네이버 지도 SDK를 동적 로드 후 지도 생성.",
        "식당 선택: 사이드바(PC) 또는 하단 패널(모바일) 버튼, 또는 지도 마커 클릭.",
        "경로 표시: T-map 보행 API 호출 → GeoJSON 파싱 → Polyline 그리기 → fitBounds()로 화면 맞춤.",
    ]
    for i, s in enumerate(steps, 1):
        add_bullet(doc, f"{i}. {s}")

    # 3. 파일별 평가
    add_heading(doc, "3. 파일별 상세 평가", 1)

    files_eval = [
        (
            "3.1 mapUtils.js — 공통 비즈니스 로직 (★★★★☆)",
            "식당 좌표, GPS 옵션, 줌 레벨 등 상수를 중앙 관리한다. addMyLocationMarker, addCafeteriaMarkers, fetchPedestrianRoute 함수로 PC/모바일에서 재사용한다. InfoWindow HTML은 CSS 클래스(map-info-window)와 연동하여 스타일을 분리하였다.",
            [
                "잘된 점: 상수 중앙 관리, 재사용 가능한 함수, GeoJSON 파싱 간결화",
                "개선 여지: T-map API 키 하드코딩 → 환경 변수 분리 필요",
                "개선 여지: API 실패 시 사용자에게 보이는 에러 메시지 추가",
            ],
        ),
        (
            "3.2 Map.jsx — 데스크톱 UI (★★★★☆)",
            "좌측 30% 사이드바(로고 + 식당 선택)와 우측 70% 지도 영역으로 구성된다. useEffect를 GPS 조회, SDK 로드, 지도 초기화, 리사이즈 4단계로 분리하였다.",
            [
                "잘된 점: openCafeteriaRoute()로 마커·버튼 클릭 로직 통합",
                "잘된 점: useRef로 지도 인스턴스 관리, cleanup으로 메모리 누수 방지",
                "잘된 점: 로딩/에러 상태 UI, 「내 위치로 이동」 버튼 제공",
                "개선 여지: mobile_map.jsx와 중복 코드 → 커스텀 Hook 추출 가능",
            ],
        ),
        (
            "3.3 mobile_map.jsx — 모바일 UI (★★★★☆)",
            "전체 화면 지도 + 하단 플로팅 패널 구조이다. mapUtils.js를 공유하여 데스크톱과 동일한 지도·경로 로직을 사용한다.",
            [
                "잘된 점: MOBILE_FIT_BOUNDS_PADDING(bottom: 240)으로 하단 탭에 가려지지 않게 경로 표시",
                "잘된 점: Map_style.css import로 마커·InfoWindow 스타일 공유",
                "개선 여지: 「내 위치로 이동」 버튼 미구현 (PC 대비 기능 비대칭)",
            ],
        ),
        (
            "3.4 Map_style.css — 데스크톱 + 지도 오버레이 (★★★★★)",
            "Flexbox 30:70 레이아웃, BEM 유사 네이밍, 선택 상태 시각 피드백(그라데이션 + accent bar), #3b82f6 블루 컬러 시스템 통일.",
            [
                "잘된 점: hover/active transition, 인라인 스타일 제거 후 CSS 클래스화",
                "개선 여지: CSS 변수 도입 시 테마 변경 용이",
            ],
        ),
        (
            "3.5 mobile_map_style.css — 모바일 전용 (★★★★☆)",
            "position: fixed 풀스크린, backdrop-filter 블러, border-radius: 20px, safe-area-inset-bottom으로 iPhone 노치 대응.",
            [
                "잘된 점: 4열 그리드 + 360px 이하 2열 전환",
                "잘된 점: mobile- 접두사로 PC CSS와 네임스페이스 분리",
                "개선 여지: lightskyblue 테두리 → PC와 동일 블루 계열 통일",
            ],
        ),
    ]

    for title, desc, bullets in files_eval:
        add_heading(doc, title, 2)
        add_para(doc, desc)
        for b in bullets:
            add_bullet(doc, b)

    # 4. 기술적 특징
    add_heading(doc, "4. 기술적 특징 (발표 강조 포인트)", 1)
    add_heading(doc, "4.1 반응형 설계", 2)
    add_para(doc, "index.js에서 matchMedia와 change 이벤트로 768px breakpoint를 적용한다. CSS @media만이 아니라 컴포넌트 레벨에서 PC/모바일 UI 구조 자체를 분리하였다.")

    add_heading(doc, "4.2 관심사 분리", 2)
    add_table(
        doc,
        ["계층", "담당 파일"],
        [
            ["UI·상태", "Map.jsx, mobile_map.jsx"],
            ["지도·API 로직", "mapUtils.js"],
            ["스타일", "Map_style.css, mobile_map_style.css"],
            ["진입점·분기", "index.js"],
        ],
    )

    add_heading(doc, "4.3 Imperative + Declarative 혼합", 2)
    add_para(
        doc,
        "React는 선언적 UI(버튼, 사이드바)를, 네이버 지도 SDK는 명령형 API(Marker, Polyline, fitBounds)를 사용한다. "
        "useRef로 지도 인스턴스를 직접 제어하는 것은 지도 SDK 연동 프로젝트에서 흔히 쓰는 실용적 패턴이다.",
    )

    add_heading(doc, "4.4 이중 API 연동", 2)
    add_para(doc, "네이버 지도는 한국 지도 타일·좌표계에 최적화되어 표시를, T-map은 보행 경로·거리·시간 데이터를 담당한다. 「지도 표시」와 「경로 계산」을 역할별로 분리한 설계이다.")

    # 5. 종합 평가
    add_heading(doc, "5. 종합 평가", 1)
    add_table(
        doc,
        ["평가 항목", "점수", "코멘트"],
        [
            ["기능 완성도", "★★★★☆", "GPS, 지도, 마커, 경로, 시간/거리 표시 구현"],
            ["코드 구조", "★★★★☆", "mapUtils 분리, Hooks 활용 양호"],
            ["UI/UX", "★★★★☆", "PC 사이드바·모바일 플로팅 패널, 선택 상태, 로딩 UI"],
            ["반응형", "★★★★☆", "breakpoint + 별도 컴포넌트"],
            ["유지보수성", "★★★☆☆", "Map/mobile 중복, API 키 하드코딩"],
            ["보안", "★★☆☆☆", "T-map 키 노출, 네이버 키는 env 처리됨"],
        ],
    )
    add_para(
        doc,
        "총평: 실제 외부 API 2종 + GPS + 반응형 UI를 연동한 수준으로, 웹 프론트엔드 팀 프로젝트로 충분히 발표·보고 가능한 완성도이다. "
        "mapUtils.js로 공통 로직을 추출한 점, 모바일 fitBounds 패딩 조정 등 사용자 경험을 코드로 반영한 부분이 설계 의도를 설명하기 좋다.",
        bold=False,
    )

    doc.add_page_break()

    # 6. 보고서용 문단
    add_heading(doc, "6. 보고서·발표용 문단 (복사 가능)", 1)

    sections = [
        (
            "6.1 연구/개발 배경",
            "계명대학교 캠퍼스에는 여러 학식당이 분산되어 있어, 신입생·교환학생 등에게 「어느 식당이 가까운지」 안내하는 도구가 필요하다. "
            "기존 지도 앱은 학교 식당 특화 정보가 부족하므로, 캠퍼스 학식당 전용 지도 서비스를 React 웹 앱으로 구현하였다.",
        ),
        (
            "6.2 구현 방법",
            "React 함수형 컴포넌트와 Hooks를 사용하여 UI와 상태를 관리하고, 네이버 지도 JavaScript API로 지도를 렌더링하였다. "
            "사용자 위치는 W3C Geolocation API로 획득하며, 식당까지의 보행 경로는 SK T-map Open API의 Pedestrian Route API를 호출하여 "
            "GeoJSON 형태로 받아 지도 위 Polyline으로 시각화하였다. 화면 너비 768px을 기준으로 PC용 사이드바 레이아웃과 "
            "모바일용 하단 플로팅 패널 레이아웃을 분리하여 반응형을 구현하였다.",
        ),
        (
            "6.3 기대 효과",
            "사용자는 별도 앱 설치 없이 웹 브라우저만으로 현재 위치 기준 학식당 경로를 확인할 수 있다. "
            "소요 시간과 거리 정보를 함께 제공하여 점심 시간 등 제한된 시간 내 식당 선택을 지원한다.",
        ),
        (
            "6.4 한계 및 향후 개선",
            "(1) T-map API 키를 환경 변수로 분리하고 백엔드 프록시를 도입하여 키 노출 방지\n"
            "(2) useNaverMap 커스텀 Hook으로 PC/모바일 중복 코드 제거\n"
            "(3) API 실패 시 사용자-facing 에러 메시지 추가\n"
            "(4) 식당 메뉴·운영 시간 등 메타데이터 연동\n"
            "(5) PWA 또는 Capacitor로 네이티브 앱 패키징",
        ),
    ]
    for title, body in sections:
        add_heading(doc, title, 2)
        add_para(doc, body)

    # 7. 발표 구성
    add_heading(doc, "7. 발표 슬라이드 구성 제안 (5~7분)", 1)
    slides = [
        "문제 정의 — 캠퍼스 학식당 찾기 어려움",
        "솔루션 — GPS + 지도 + 보행 경로",
        "데모 — PC / 모바일 화면 시연",
        "아키텍처 — 컴포넌트·API 구조도",
        "기술 스택 — React, Naver Maps, T-map, Geolocation",
        "코드 구조 — mapUtils 분리, 반응형 분기",
        "한계 & 개선 — API 키, 에러 처리, 기능 확장",
    ]
    for i, s in enumerate(slides, 1):
        add_bullet(doc, f"슬라이드 {i}: {s}")

    add_heading(doc, "8. 데모 시연 시나리오", 1)
    demo = [
        "브라우저에서 localhost:3000 접속",
        "「GPS 확인 중…」→ 지도 로드 확인",
        "「공대 식당」 클릭 → 파란 경로선 + 소요시간/거리 InfoWindow 확인",
        "개발자 도구 모바일 뷰(375px) 전환 → 하단 탭 UI 확인",
        "(PC) 「내 위치로 이동」 클릭 → 지도 중심 복귀",
    ]
    for i, s in enumerate(demo, 1):
        add_bullet(doc, f"{i}. {s}")

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    build()
